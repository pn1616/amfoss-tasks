import os
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, ContextTypes, MessageHandler, filters, CallbackQueryHandler, ConversationHandler
from googleapiclient.discovery import build
from docx import Document
import pandas as pd

AWAITING_GENRE, AWAITING_BOOK_NAME, AWAITING_ACTION, AWAITING_BOOK_TWO = range(4)

tokenn = '6324389810:AAGU2lsXHWjQbDr1lytAL9o0KVpSQ4Tl-ig'
google_api = 'AIzaSyC03ktoLpKn_QfVJyoA_fyN1Qawcoq_vjY'
books_service = build('books', 'v1', developerKey = google_api)

user_states = {}
reading_list = []

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE)-> None:
    await update.message.reply_text("hey welcome! this is a library telegram bot! type /help to know all the commands")

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    help_text = (
        "/start - welcome message\n" 
        "/book - type in a genre and get a csv file of top books\n"
        "/preview - type the book and get a preview link\n"
        "/list - allows you to make you're own reading list\n"
        "/help - displays this message"
    )
    await update.message.reply_text(help_text)
    
async def book(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_states[update.message.from_user.id] = 'awaiting_genre'
    await update.message.reply_text("enter the genre you want")
    return AWAITING_GENRE

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.message.from_user.id
    state = user_states.get(user_id)   
    
    if state== 'awaiting_genre':
        genre  = update.message.text
        results = books_service.volumes().list(q=f'subject:{genre}', maxResults=20).execute()
        books = results.get('items', [])
        
        if books:
            await update.message.reply_text(f"here you go, {genre} it is")
            data = []
            for book in books:
               title = book['volumeInfo'].get('title', 'No Title')
               authors = ', '.join(book['volumeInfo'].get('authors', ['Unknown Author']))
               language = book['volumeInfo'].get('language', ['Unknown Language'])
               published_date = book['volumeInfo'].get('publishedDate', 'Unknown Date')
               description = book['volumeInfo'].get('description', 'No Description')
                
               data.append({
                    'Title': title,
                    'Authors': authors,
                    'Language' : language,
                    'Published Date': published_date,
                    'Description': description
                })
            df = pd.DataFrame(data)
            file_path = 'books_list.csv'
            df.to_csv(file_path, index=False)  
            
            with open(file_path,'rb') as file:
                await update.message.reply_document(document=file, filename=file_path)
            os.remove(file_path)
        else:
            await update.message.reply_text("oopsie daisy")
        user_states.pop(user_id, None)
        return ConversationHandler.END
   
    else:
        await update.message.reply_text("Uhmm you suck")
        return ConversationHandler.END
     
        
async def preview(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_states[update.message.from_user.id] = 'awaiting_book_name'
    await update.message.reply_text("enter the book to read the first few pages")
    return AWAITING_BOOK_NAME
    
async def handle_book_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user_id = update.message.from_user.id
    state = user_states.get(user_id)
    
    if state == 'awaiting_book_name':
        book_name = update.message.text
        results = books_service.volumes().list(q=book_name, maxResults=1).execute()
        books = results.get('items', [])
        
        if books:
            book = books[0]
            preview_link = book['volumeInfo'].get('previewLink', 'No preview available')
            await update.message.reply_text(f"Preview link for '{book_name}': {preview_link}")
        else:
            await update.message.reply_text("did you dream up this book?")
        
        user_states.pop(user_id, None)
        return ConversationHandler.END
    else:
        await update.message.reply_text("error")   
        return ConversationHandler.END
    
async def list_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    user_states[update.message.from_user.id] = 'awaiting_book_two'
    await update.message.reply_text("enter the book title for the reading list")
    return AWAITING_BOOK_TWO

async def reading_list_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    keyboard = [
        [InlineKeyboardButton("add a book", callback_data='add')],
        [InlineKeyboardButton("delete a book", callback_data='delete')],
        [InlineKeyboardButton('view reading list', callback_data='view')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("choose the following options:", reply_markup=reply_markup)
    return AWAITING_ACTION

async def handle_button_press(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    
    if query.data == 'add':
        book_name = user_states.get(user_id)
        if book_name: 
            reading_list.append(book_name)
            await query.edit_message_text(f"added {book_name} to your reading list")
        else:
            await query.edit_message_text("no book name was provided")
            
    elif query.data == 'delete':
        if reading_list:
            book_name = user_states.get(user_id)
            if book_name and book_name in reading_list:
                reading_list.remove(book_name)
                await query.edit_message_text(f"deleted successfully")
            else:
                await query.edit_message_text("book wasnt found ")
        else:
            await query.edit_message_text("no books in reading list")
            
    elif query.data =='view':
        if reading_list:
            document = Document()
            document.add_heading('reading list', 0)
            for book in reading_list:
                document.add_paragraph(book)
            file_path = 'reading_list.docx'
            document.save(file_path)
            
            with open(file_path, 'rb') as file:
                await context.bot.send_document(chat_id=user_id, document=file, filename=file_path)
        else:
            await query.edit_message_text("empty")
    return ConversationHandler.END
            
    
if __name__ == '__main__':
    print("Starting bot.. ")
    app = Application.builder().token(tokenn).build()
    
    conv_handler = ConversationHandler(
        entry_points=[
            CommandHandler('book', book),
            CommandHandler('preview', preview),
            CommandHandler('list', list_command)
        ],
        states={
            AWAITING_GENRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text)],
            AWAITING_BOOK_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_book_name)],
            AWAITING_BOOK_TWO: [MessageHandler(filters.TEXT & ~filters.COMMAND, reading_list_command )],
            AWAITING_ACTION: [CallbackQueryHandler(handle_button_press)]
        },
        fallbacks=[CommandHandler('start', start)])

    
    app.add_handler(conv_handler)
    app.add_handler(CommandHandler('start', start))
    app.add_handler(CommandHandler('help', help_command))
    
   
    app.run_polling()
    